import os
from typing import Optional
from docx import Document
from striprtf.striprtf import rtf_to_text
import markdown
import chardet
import mimetypes
import json

class FileReader:
    def __init__(self):
        # Initialize supported file types
        self.text_extensions = {
            '.txt', '.md', '.markdown', '.py', '.js', '.html', '.css', '.json', 
            '.xml', '.yaml', '.yml', '.ini', '.cfg', '.conf', '.log', '.csv',
            '.rtf', '.tex', '.sh', '.bat', '.ps1', '.r', '.sql'
        }
        
    def detect_encoding(self, file_path: str) -> str:
        """Detect the file encoding."""
        with open(file_path, 'rb') as f:
            raw = f.read()
            result = chardet.detect(raw)
            return result['encoding'] or 'utf-8'

    def read_file(self, file_path: str) -> Optional[str]:
        """
        Reads and returns the contents of a file at the given path.
        
        Args:
            file_path (str): Path to the file to read
            
        Returns:
            Optional[str]: Contents of the file if successful, None if file cannot be read
        """
        try:
            if not os.path.exists(file_path):
                return f"Error: File {file_path} does not exist"
                
            extension = os.path.splitext(file_path)[1].lower()
            mime_type, _ = mimetypes.guess_type(file_path)
            
            # Handle different file types
            if extension == '.docx':
                doc = Document(file_path)
                text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
                for table in doc.tables:
                    for row in table.rows:
                        text += '\n' + ' | '.join([cell.text for cell in row.cells])
                return text
                
            elif extension == '.rtf':
                with open(file_path, 'r', encoding='utf-8') as f:
                    rtf_text = f.read()
                return rtf_to_text(rtf_text)
                
            elif extension in ['.md', '.markdown']:
                encoding = self.detect_encoding(file_path)
                with open(file_path, 'r', encoding=encoding) as f:
                    md_text = f.read()
                    # Convert Markdown to plain text while preserving structure
                    html = markdown.markdown(md_text)
                    # Could use BeautifulSoup here to better format HTML if needed
                    return html
                    
            elif extension == '.pdf':
                return "PDF parsing not implemented yet"
                
            elif extension in self.text_extensions or (mime_type and mime_type.startswith('text/')):
                # Handle all text-based files with encoding detection
                encoding = self.detect_encoding(file_path)
                with open(file_path, 'r', encoding=encoding) as f:
                    return f.read()
                    
            else:
                return f"Error: Unsupported file type {extension}"
                
        except UnicodeDecodeError as e:
            return f"Error: File encoding issue - {str(e)}"
        except Exception as e:
            return f"Error reading file: {str(e)}"

    def format_file_analysis(self, file_info: dict) -> str:
        """
        Formats the file analysis results into a readable string.
        
        Args:
            file_info (dict): The analysis results dictionary
            
        Returns:
            str: Formatted analysis string
        """
        if "error" in file_info:
            return f"Error analyzing file: {file_info['error']}"
        
        # Format file size to human-readable format
        def format_size(size_bytes):
            for unit in ['B', 'KB', 'MB', 'GB']:
                if size_bytes < 1024:
                    return f"{size_bytes:.2f} {unit}"
                size_bytes /= 1024
            return f"{size_bytes:.2f} TB"
        
        # Format timestamp to readable date
        from datetime import datetime
        last_modified = datetime.fromtimestamp(file_info["last_modified"]).strftime("%Y-%m-%d %H:%M:%S")
        
        # Build the formatted output
        output = [
            "📄 File Analysis Report",
            "─" * 30,
            f"📌 Basic Information:",
            f"• Type: {file_info['file_type']}",
            f"• MIME: {file_info['mime_type']}",
            f"• Size: {format_size(file_info['size'])}",
            f"• Last Modified: {last_modified}",
            f"• Encoding: {file_info['encoding']}",
            "",
            "📊 Content Summary:"
        ]
        
        # Add file-type specific information
        summary = file_info["summary"]
        if file_info["file_type"] == ".docx":
            output.extend([
                f"• Paragraphs: {summary['paragraphs']}",
                f"• Tables: {summary['tables']}",
                f"• Sections: {summary['sections']}",
                f"• Lines: {summary['line_count']}",
                f"• Characters: {summary['char_count']:,}"
            ])
        elif file_info["file_type"] in [".md", ".markdown"]:
            output.extend([
                f"• Headers: {summary['headers']}",
                f"• Links: {summary['links']}",
                f"• Code Blocks: {summary['code_blocks']}",
                f"• Lines: {summary['line_count']}",
                f"• Characters: {summary['char_count']:,}"
            ])
        else:
            output.extend([
                f"• Lines: {summary['line_count']}",
                f"• Words: {summary['word_count']:,}",
                f"• Characters: {summary['char_count']:,}"
            ])
        
        # Add preview of content if available
        if file_info["content"]:
            preview_length = 200
            content_preview = file_info["content"][:preview_length]
            if len(file_info["content"]) > preview_length:
                content_preview += "..."
            output.extend([
                "",
                "📝 Content Preview:",
                "─" * 30,
                content_preview
            ])
        
        return "\n".join(output)

    def analyze_file(self, file_path: str) -> dict:
        """
        Analyzes a file and returns key information about it.
        
        Args:
            file_path (str): Path to the file to analyze
            
        Returns:
            dict: Analysis results including:
                - file_type: Type/extension of the file
                - mime_type: MIME type of the file
                - size: Size in bytes
                - encoding: Detected file encoding
                - content: File contents
                - summary: File-type specific statistics
        """
        try:
            extension = os.path.splitext(file_path)[1].lower()
            mime_type, _ = mimetypes.guess_type(file_path)
            content = self.read_file(file_path)
            
            # Basic file info
            file_info = {
                "file_type": extension,
                "mime_type": mime_type,
                "size": os.path.getsize(file_path),
                "encoding": self.detect_encoding(file_path),
                "content": content,
                "last_modified": os.path.getmtime(file_path)
            }
            
            # File type specific analysis
            if extension == '.docx':
                doc = Document(file_path)
                file_info["summary"] = {
                    "paragraphs": len(doc.paragraphs),
                    "tables": len(doc.tables),
                    "sections": len(doc.sections),
                    "line_count": len(content.split('\n')),
                    "char_count": len(content)
                }
            elif extension in ['.md', '.markdown']:
                file_info["summary"] = {
                    "line_count": len(content.split('\n')),
                    "char_count": len(content),
                    "headers": content.count('#'),  # Approximate header count
                    "links": content.count(']('),   # Approximate link count
                    "code_blocks": content.count('```')  # Approximate code block count
                }
            else:
                file_info["summary"] = {
                    "line_count": len(content.split('\n')) if content else 0,
                    "char_count": len(content) if content else 0,
                    "word_count": len(content.split()) if content else 0
                }
            
            return {
                "result": file_info,
                "formatted_output": self.format_file_analysis(file_info)
            }
            
        except Exception as e:
            return {
                "error": f"Failed to analyze file: {str(e)}",
                "formatted_output": f"Error: Failed to analyze file - {str(e)}"
            } 